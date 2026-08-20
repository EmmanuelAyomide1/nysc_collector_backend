import io

from apps.payments.models import Payment
from apps.users.models import CustomUser


def get_members_who_paid(payment_item):
    return (
        Payment.objects.filter(
            payment_item=payment_item, status=Payment.Status.SUCCESSFUL
        )
        .select_related("member")
        .order_by("member__last_name", "member__first_name")
    )


def group_payments_by_batch(payments):
    """Group payments by member batch, in Batch-enum order, dropping empty batches."""
    by_batch = {}
    for payment in payments:
        by_batch.setdefault(payment.member.batch, []).append(payment)

    return [
        (label, by_batch[value])
        for value, label in CustomUser.Batch.choices
        if value in by_batch
    ]


def export_payment_item_report_pdf(payment_item):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    payments = list(get_members_who_paid(payment_item))
    batches = group_payments_by_batch(payments)
    styles = getSampleStyleSheet()

    table_style = TableStyle(
        [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#111827")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            (
                "ROWBACKGROUNDS",
                (0, 1),
                (-1, -1),
                [colors.white, colors.HexColor("#f3f4f6")],
            ),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]
    )

    elements = [
        Paragraph(f"Payment Report: {payment_item.name}", styles["Title"]),
        Paragraph(f"Members who have paid ({len(payments)})", styles["Normal"]),
    ]

    for batch_index, (batch_label, batch_payments) in enumerate(batches):
        if batch_index > 0:
            elements.append(PageBreak())
        elements.append(Spacer(1, 0.5 * cm))
        elements.append(
            Paragraph(f"{batch_label} ({len(batch_payments)})", styles["Heading2"])
        )
        elements.append(Spacer(1, 0.3 * cm))

        data = [["S/N", "Name", "State Code"]]
        for index, payment in enumerate(batch_payments, start=1):
            member = payment.member
            data.append(
                [
                    str(index),
                    f"{member.first_name} {member.last_name}",
                    member.state_code,
                ]
            )

        table = Table(data, colWidths=[1.5 * cm, 8 * cm, 6 * cm])
        table.setStyle(table_style)
        elements.append(table)

    buffer = io.BytesIO()
    SimpleDocTemplate(buffer, pagesize=A4).build(elements)
    buffer.seek(0)
    return buffer


def export_payment_item_report_docx(payment_item):
    from docx import Document

    payments = list(get_members_who_paid(payment_item))
    batches = group_payments_by_batch(payments)

    document = Document()
    document.add_heading(f"Payment Report: {payment_item.name}", level=1)
    document.add_paragraph(f"Members who have paid ({len(payments)})")

    for batch_index, (batch_label, batch_payments) in enumerate(batches):
        if batch_index > 0:
            document.add_page_break()
        document.add_heading(f"{batch_label} ({len(batch_payments)})", level=2)

        table = document.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        header_cells = table.rows[0].cells
        header_cells[0].text = "S/N"
        header_cells[1].text = "Name"
        header_cells[2].text = "State Code"

        for index, payment in enumerate(batch_payments, start=1):
            member = payment.member
            row_cells = table.add_row().cells
            row_cells[0].text = str(index)
            row_cells[1].text = f"{member.first_name} {member.last_name}"
            row_cells[2].text = member.state_code

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
